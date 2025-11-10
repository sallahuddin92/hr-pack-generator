# ===================================================================
# 🇲🇾 HR Pack Generator - BACKEND API (All-in-One)
#
# VERSION 3: Upgraded logo sanitization to strip bad metadata (DPI)
#          and improved temp file cleanup.
# ===================================================================

import os
import zipfile
import tempfile
import shutil
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image as PILImage # <-- Import PIL's Image

# --- 1. Initialize the Flask App ---
app = Flask(__name__)

# --- 2. Enable CORS ---
CORS(app)

# --- 3. UPGRADED HELPER: Logo Sanitizer ---
# This function fixes logos that are "too large" or have bad metadata/DPI
def sanitize_logo(logo_upload_file, base_dir):
    """
    Opens, rebuilds, and saves the logo to a clean format,
    stripping all corrupted or problematic metadata.
    Returns the path to the new, sanitized logo.
    """
    try:
        # Save the uploaded file temporarily
        temp_logo_path = os.path.join(base_dir, "temp_original_logo")
        logo_upload_file.save(temp_logo_path)
        
        # Open with Pillow
        with PILImage.open(temp_logo_path) as image:
            # Convert to RGBA for consistent handling
            # This is a key step to ensure we can create a clean mask
            image = image.convert("RGBA")
            
            # Resize to a max width/height of 500px, keeping aspect ratio
            image.thumbnail((500, 500))
            
            # --- The "Rebuild" Process ---
            # 1. Create a new, blank canvas with a transparent background
            #    Use the thumbnail's *actual* new size
            new_canvas = PILImage.new("RGBA", image.size, (255, 255, 255, 0))
            
            # 2. Paste the resized image onto the blank canvas.
            #    We use the image's alpha channel (its own transparency) as the mask.
            #    This rebuilds the image from scratch and strips all bad metadata/DPI.
            new_canvas.paste(image, (0, 0), image)
            
            # 3. Save the *new canvas* as our clean file
            sanitized_logo_path = os.path.join(base_dir, "logo_sanitized.png")
            new_canvas.save(sanitized_logo_path, "PNG")
        
        # Clean up the original temp file
        os.remove(temp_logo_path)
        
        print("Logo successfully sanitized and rebuilt.")
        return sanitized_logo_path
        
    except Exception as e:
        print(f"Error sanitizing logo: {e}")
        # If sanitization fails, raise an exception
        raise Exception(f"Failed to process logo. It may be corrupted or in an unsupported format. Error: {e}")

# ===============================================================
# 4️⃣ The "Engine" (Document Generators)
#
# This is the full, un-redacted code for all generators.
# ===============================================================

# --- 4.1: Utility Functions ---

def add_footer(doc, text):
    """Adds a tagline to the footer of every page."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = text

def add_logo(doc, logo_path):
    """Adds the brand logo to the top of the document."""
    try:
        if logo_path and os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(1.3))
    except Exception as e:
        print(f"⚠️ Docx Logo add failed (is {logo_path} a valid image?): {e}")
        # Don't stop the whole process, just skip the logo

def add_bilingual_disclaimer(doc, en_text, bm_text):
    """Adds a formatted, bilingual disclaimer box."""
    p = doc.add_paragraph()
    p.add_run("EN: ").bold = True
    p.add_run(en_text).italic = True
    p.add_run("\n")
    p.add_run("BM: ").bold = True
    p.add_run(bm_text).italic = True

def add_bilingual_block(doc, en_text, bm_text):
    """Adds a standard formatted bilingual paragraph block (EN and BM)."""
    p_en = doc.add_paragraph()
    p_en.add_run("EN: ").bold = True
    p_en.add_run(en_text)
    p_bm = doc.add_paragraph()
    p_bm.add_run("BM: ").bold = True
    p_bm.add_run(bm_text)

def add_table_row(table, texts):
    """Helper to add a row to a table and populate it."""
    row = table.add_row().cells
    for i, text in enumerate(texts):
        row[i].text = text
    return row

# --- 4.2: Hiring Kit Generator ---

def make_hiring_kit(COMPANY_DETAILS, BRAND_TAGLINE, logo_path, save_folder):
    """Generates all 6 documents for the Hiring & Onboarding Kit."""
    print("Generating Hiring Kit...")
    
    # === 01_Job_Description_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Job Description Template / Templat Penerangan Kerja", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    table = doc.add_table(rows=4, cols=2); table.style = 'Table Grid'
    table.cell(0, 0).text = "Jawatan / Job Title:"; table.cell(0, 1).text = "<<Jawatan / Position>>"
    table.cell(1, 0).text = "Jabatan / Department:"; table.cell(1, 1).text = "<<Jabatan / Department>>"
    table.cell(2, 0).text = "Melapor Kepada / Reports To:"; table.cell(2, 1).text = "<<Jawatan Pengurus / Manager's Title>>"
    table.cell(3, 0).text = "Julat Gaji / Salary Range:"; table.cell(3, 1).text = "<<RM XXXX - RM XXXX>> (Anggaran) / (Estimated)"
    doc.add_paragraph(); doc.add_heading("TUJUAN JAWATAN / JOB PURPOSE", 1)
    add_bilingual_block(doc,
        f"(A brief 1-2 sentence summary of the main purpose of this role and why it exists.)\nExample: To manage all digital marketing channels for {COMPANY_DETAILS['name']}, including social media, email marketing, and SEO, to generate leads and build brand awareness.",
        f"(Ringkasan 1-2 ayat tentang tujuan utama peranan ini dan mengapa ia wujud.)\nContoh: Mengurus semua saluran pemasaran digital untuk {COMPANY_DETAILS['name']}, termasuk media sosial, pemasaran e-mel, dan SEO, untuk menjana petunjuk jualan (leads) dan membina kesedaran jenama."
    )
    doc.add_heading("TANGGUNGJAWAB UTAMA / KEY RESPONSIBILITIES", 1)
    doc.add_paragraph("(Sila senaraikan 5-8 tanggungjawab utama. / Please list 5-8 primary responsibilities.)")
    add_bilingual_block(doc,"<< Responsibility 1 (e.g., Plan and execute all social media campaigns.) >>", "<< Tanggungjawab 1 (cth., Merancang dan melaksanakan semua kempen media sosial.) >>")
    add_bilingual_block(doc,"<< Responsibility 2 (e.g., Monitor, analyze, and report on campaign performance.) >>", "<< Tanggungjawab 2 (cth., Memantau, menganalisis, dan melaporkan prestasi kempen.) >>")
    add_bilingual_block(doc,"<< Responsibility 3 (e.g., Liaise with the Sales team to align strategies.) >>", "<< Tanggungjawab 3 (cth., Berhubung dengan pasukan Jualan untuk menyelaraskan strategi.) >>")
    add_bilingual_block(doc,"<< Responsibility 4 >>", "<< Tanggungjawab 4 >>")
    doc.add_heading("KELAYAKAN & KEMAHIRAN DIPERLUKAN / QUALIFICATIONS & SKILLS REQUIRED", 1)
    add_bilingual_block(doc,"<< Qualification 1 (e.g., Diploma/Degree in Marketing...) >>", "<< Kelayakan 1 (cth., Diploma/Ijazah dalam Pemasaran...) >>")
    add_bilingual_block(doc,"<< Skill 2 (e.g., Minimum 2 years experience...) >>", "<< Kemahiran 2 (cth., Pengalaman minimum 2 tahun...) >>")
    add_bilingual_block(doc,"<< Skill 3 (e.g., Proficient in Meta Business Suite...) >>", "<< Kemahiran 3 (cth., Mahir dalam Meta Business Suite...) >>")
    add_bilingual_block(doc,"<< Skill 4 (e.g., Excellent communication skills...) >>", "<< Kemahiran 4 (cth., Kemahiran komunikasi cemerlang...) >>")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "01_Job_Description_Template.docx"))

    # === 02_Hiring_Process_Checklist.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Hiring Process Checklist / Senarai Semak Proses Pengambilan", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_paragraph("Jawatan / Position: <<Jawatan / Position>>")
    doc.add_paragraph("Pengurus Pengambilan / Hiring Manager: <<Nama Pengurus / Manager's Name>>")
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    add_table_row(table, ["Fasa / Phase", "Tugasan / Task", "Status (Tandakan / Tick)", "Nota / Notes"])
    add_table_row(table, ["1. Permohonan / Requisition", "EN: Job Description (JD) finalized & approved.\nBM: Huraian Tugas (JD) dimuktamadkan & diluluskan.", "[]", ""])
    add_table_row(table, ["", "EN: Job vacancy advertised (e.g., JobStreet, LinkedIn).\nBM: Jawatan kosong diiklankan (cth., JobStreet, LinkedIn).", "[]", "Tarikh Iklan / Date Advertised:"])
    add_table_row(table, ["2. Saringan / Screening", "EN: Applications closing date.\nBM: Tarikh tutup permohonan.", "[]", "Tarikh / Date:"])
    add_table_row(table, ["", "EN: Review applications & create shortlist (Top 10).\nBM: Semak permohonan & buat senarai pendek (10 Terbaik).", "[]", ""])
    add_table_row(table, ["3. Temu Duga / Interview", "EN: Initial phone screen (if applicable).\nBM: Saringan telefon awal (jika perlu).", "[]", ""])
    add_table_row(table, ["", "EN: Schedule Round 1 Interviews (Hiring Manager).\nBM: Jadualkan Temu Duga Pusingan 1 (Pengurus Pengambilan).", "[]", ""])
    add_table_row(table, ["", "EN: Schedule Round 2 Interviews (e.g., Technical test / Management).\nBM: Jadualkan Temu Duga Pusingan 2 (cth., Ujian teknikal / Pengurusan).", "[]", ""])
    add_table_row(table, ["", "EN: Select final candidate.\nBM: Pilih calon akhir.", "[]", "Calon / Candidate:"])
    add_table_row(table, ["", "EN: Conduct reference checks (2x).\nBM: Buat semakan rujukan (2x).", "[]", ""])
    add_table_row(table, ["4. Tawaran / Offer", "EN: Verbal offer made and accepted.\nBM: Tawaran lisan dibuat dan diterima.", "[]", ""])
    add_table_row(table, ["", "EN: Formal Letter of Offer (Contract) prepared & sent.\nBM: Surat Tawaran Rasmi (Kontrak) disediakan & dihantar.", "[]", ""])
    add_table_row(table, ["", "EN: Signed Letter of Offer received.\nBM: Surat Tawaran yang ditandatangani diterima.", "[]", ""])
    add_table_row(table, ["5. Pra-Kemasukan / Pre-boarding", "EN: Collect new hire documents (IC, bank, education certs).\nBM: Kumpul dokumen pekerja baharu (IC, bank, sijil pendidikan).", "[]", ""])
    add_table_row(table, ["", "EN: Register for EPF, SOCSO, EIS.\nBM: Daftar untuk KWSP, PERKESO, EIS.", "[]", ""])
    add_table_row(table, ["", "EN: Send Welcome Letter.\nBM: Hantar Surat Aluan.", "[]", ""])
    add_table_row(table, ["", "EN: Announce new hire to the team.\nBM: Umumkan pekerja baharu kepada pasukan.", "[]", ""])
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "02_Hiring_Process_Checklist.docx"))

    # === 03_Candidate_Interview_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Candidate Interview Template / Templat Temuduga Calon", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    table_info = doc.add_table(rows=2, cols=2)
    table_info.cell(0, 0).text = "Nama Calon / Candidate Name:"; table_info.cell(0, 1).text = "Jawatan / Position:"
    table_info.cell(1, 0).text = "Penemu Duga / Interviewer:"; table_info.cell(1, 1).text = "Tarikh / Date:"
    doc.add_paragraph()
    doc.add_paragraph("Skala Penilaian / Rating Scale:", style='Intense Quote')
    doc.add_paragraph("1 = Lemah / Poor\n2 = Sederhana / Fair\n3 = Baik / Good\n4 = Sangat Baik / Very Good\n5 = Cemerlang / Excellent")
    doc.add_paragraph()
    table_main = doc.add_table(rows=1, cols=3); table_main.style = 'Table Grid'
    add_table_row(table_main, ["Kompetensi / Competency\nSoalan Dicadangkan / Suggested Question", "Penilaian (1-5) / Rating (1-5)", "Nota Penemu Duga / Interviewer's Notes"])
    add_table_row(table_main, ["Pengenalan / Introduction\nEN: Tell me about yourself...\nBM: Ceritakan tentang diri anda...", "N/A", ""])
    add_table_row(table_main, ["Kemahiran Teknikal / Technical Skills\nEN: This role requires X...\nBM: Jawatan ini memerlukan X...", "", ""])
    add_table_row(table_main, ["Penyelesaian Masalah / Problem Solving\nEN: Tell me about a time you faced a difficult challenge...\nBM: Ceritakan satu masa anda menghadapi cabaran sukar...", "", ""])
    add_table_row(table_main, ["Kerja Berpasukan / Teamwork\nEN: How do you handle disagreements with a colleague?\nBM: Bagaimana anda mengendalikan perselisihan faham...", "", ""])
    add_table_row(table_main, ["Inisiatif / Initiative\nEN: Describe a project or idea you started...\nBM: Terangkan satu projek atau idea yang anda mulakan...", "", ""])
    add_table_row(table_main, [f"Kesesuaian Budaya / Culture Fit\nEN: Our company values X...\nBM: Syarikat kami ({COMPANY_DETAILS['name']}) menghargai X...", "", ""])
    add_table_row(table_main, ["Soalan Calon / Candidate's Questions\nEN: Do you have any questions for me/us?\nBM: Adakah anda mempunyai soalan...", "N/A", "(Nota: Adakah soalan calon bernas?)"])
    doc.add_paragraph()
    doc.add_heading("Rumusan Penemu Duga / Interviewer's Summary", 1)
    doc.add_paragraph("Kekuatan / Strengths:\n<<...>>\n")
    doc.add_paragraph("Kelemahan / Weaknesses:\n<<...>>\n")
    doc.add_paragraph("Syor Pengambilan / Hiring Recommendation:\n[ ] Syor Ambil / Recommend to Hire\n[ ] Pertimbangkan / Consider\n[ ] Jangan Ambil / Do Not Hire")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "03_Candidate_Interview_Template.docx"))

    # === 04_Letter_of_Offer_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Surat Tawaran Pelantikan / Letter of Offer of Employment", 0)
    add_bilingual_disclaimer(doc, "Once signed, this document may constitute a legally binding employment contract.", "Setelah ditandatangani, dokumen ini boleh menjadi satu kontrak pekerjaan yang sah di sisi undang-undang.")
    doc.add_paragraph("Tarikh / Date: <<Tarikh / Date>>\n\n")
    doc.add_paragraph("KEPADA / TO:\n<<Nama Penuh Calon / Candidate Full Name>>\n(No. K/P/NRIC No.: <<NRIC No.>>)\n<<Alamat Calon / Candidate Address>>\n\n")
    doc.add_heading("PERKARA: SURAT TAWARAN PELANTIKAN / RE: LETTER OF OFFER OF EMPLOYMENT", 1)
    add_bilingual_block(doc, f"We are pleased to offer you, <<Candidate Name>>, (\"the Employee\") employment with {COMPANY_DETAILS['name']} (\"the Company\")...", f"Dengan sukacitanya, {COMPANY_DETAILS['name']} (\"Syarikat\") ingin menawarkan anda pelantikan...")
    doc.add_heading("1.0 Jawatan / Position", 2)
    doc.add_paragraph("Jawatan / Title: <<Jawatan / Position>>\nTarikh Mula / Start Date: <<Tarikh Mula / Start Date>>\nMelapor Kepada / Reports To: <<Jawatan Pengurus / Manager's Title>>")
    doc.add_heading("2.0 Tempoh Percubaan / Probationary Period", 2)
    add_bilingual_block(doc, "You will serve a probationary period of <<Three (3)>> months... Your confirmation is subject to your satisfactory performance.", "Anda akan berkhidmat dalam tempoh percubaan selama <<Tiga (3)>> bulan... Pengesahan jawatan anda adalah tertakluk kepada prestasi yang memuaskan.")
    doc.add_heading("3.0 Gaji & Elaun / Salary & Allowances", 2)
    doc.add_paragraph("Pakej saraan anda adalah seperti berikut: / Your remuneration package will be as follows:")
    doc.add_paragraph("• Gaji Pokok / Basic Salary: RM <<XXXX.XX>> sebulan / per month.\n• Elaun Tetap / Fixed Allowances: RM <<XXXX.XX>> sebulan / per month (jika ada / if any).")
    doc.add_heading("4.0 Potongan Berkanun / Statutory Contributions", 2)
    add_bilingual_block(doc, "Your salary will be subject to statutory deductions for EPF, SOCSO, and EIS...", "Gaji anda akan tertakluk kepada potongan berkanun untuk KWSP, PERKESO, dan EIS...")
    doc.add_heading("5.0 Waktu Bekerja / Working Hours", 2)
    add_bilingual_block(doc, f"Your normal working hours shall be 45 hours per week... (Our company's standard hours are {COMPANY_DETAILS['working_hours']}).", f"Waktu bekerja biasa anda adalah 45 jam seminggu... (Waktu standard syarikat kami ialah {COMPANY_DETAILS['working_hours']}).")
    doc.add_heading("6.0 Cuti / Leave Entitlements", 2)
    add_bilingual_block(doc, "You will be entitled to leave in accordance with the Employment Act 1955...", "Anda layak mendapat cuti selaras dengan Akta Kerja 1955...")
    doc.add_paragraph("• Cuti Tahunan / Annual Leave: <<8/12/16>> days...\n• Cuti Sakit / Sick Leave: <<14/18/22>> days...\n• Cuti Hospitalisasi / Hospitalisation Leave: 60 days...\n• Cuti Bersalin / Maternity Leave: 98 consecutive days...\n• Cuti Paterniti / Paternity Leave: 7 consecutive days...")
    doc.add_heading("7.0 Penamatan / Termination", 2)
    add_bilingual_block(doc, "During the probationary period, the notice period... is <<Two (2)>> weeks...", "Semasa tempoh percubaan, tempoh notis... adalah <<Dua (2)>> minggu...")
    doc.add_heading("8.0 Buku Panduan Pekerja / Employee Handbook", 2)
    add_bilingual_block(doc, f"This offer is subject to all terms... in the {COMPANY_DETAILS['name']} Employee Handbook...", f"Tawaran ini tertakluk kepada semua terma... dalam Buku Panduan Pekerja {COMPANY_DETAILS['name']}...")
    doc.add_heading("9.0 Tawaran Bersyarat / Conditional Offer", 2)
    add_bilingual_block(doc, "This offer is conditional upon (i) you providing satisfactory proof of your qualifications...", "Tawaran ini adalah bersyarat, tertakluk kepada (i) anda mengemukakan bukti kelayakan...")
    doc.add_paragraph()
    add_bilingual_block(doc, "If you agree to these terms, please sign... by <<Tarikh Akhir / Expiry Date>>.", "Jika anda bersetuju dengan terma-terma ini, sila tandatangani... sebelum <<Tarikh Akhir / Expiry Date>>.")
    doc.add_paragraph(f"\nYang benar, / Sincerely,\n\n___________________\n{COMPANY_DETAILS['hr_name']}\n{COMPANY_DETAILS['hr_title']}\n{COMPANY_DETAILS['name']}")
    doc.add_page_break()
    doc.add_heading("SLIP PENERIMAAN / ACCEPTANCE SLIP", 1)
    add_bilingual_block(doc, f"I, <<Candidate's Full Name>>, hereby accept the offer of employment with {COMPANY_DETAILS['name']}...", f"Saya, <<Nama Penuh Calon>>, dengan ini menerima tawaran pekerjaan dengan {COMPANY_DETAILS['name']}...")
    doc.add_paragraph("\n\n___________________\nTandatangan / Signature:\nNama / Name:\nNo. K/P/NRIC No.:\nTarikh / Date:")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "04_Letter_of_Offer_Template.docx"))

    # === 05_New_Employee_Welcome_Letter.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("New Employee Welcome Letter / Surat Aluan Pekerja Baharu", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_paragraph(f"Subjek: Selamat Datang ke Pasukan {COMPANY_DETAILS['name']}! / Subject: Welcome to the {COMPANY_DETAILS['name']} Team!")
    doc.add_heading("English Version:", 1)
    doc.add_paragraph(f"Dear <<Candidate Name>>,\n\nWelcome to the team! We are all incredibly excited to have you join {COMPANY_DETAILS['name']} as our new <<Position Title>>...\n\nHere are the details for your first day:\n• Arrival Time: <<e.g., 9:00 AM>>\n• Address: {COMPANY_DETAILS['address']}\n• Reporting To: Please ask for {COMPANY_DETAILS['hr_name']} upon arrival.\n• Dress Code: Our company dress code is {COMPANY_DETAILS['dress_code']}.\n• First Day: We have a light orientation planned...\n\nBest regards,\n{COMPANY_DETAILS['hr_name']}\n{COMPANY_DETAILS['hr_title']}")
    doc.add_heading("Bahasa Melayu Version:", 1)
    doc.add_paragraph(f"Kepada <<Nama Calon>>,\n\nSelamat datang ke pasukan! Kami semua amat teruja dengan penyertaan anda ke {COMPANY_DETAILS['name']} sebagai <<Jawatan>> baharu kami...\n\nBerikut adalah butiran untuk hari pertama anda:\n• Waktu Ketibaan: <<cth., 9:00 Pagi>>\n• Alamat: {COMPANY_DETAILS['address']}\n• Lapor Diri Kepada: Sila cari {COMPANY_DETAILS['hr_name']} semasa tiba.\n• Etika Pakaian: Etika pakaian syarikat kami ialah {COMPANY_DETAILS['dress_code']}.\n• Hari Pertama: Kami telah merancang sesi orientasi ringan...\n\nSalam hormat,\n{COMPANY_DETAILS['hr_name']}\n{COMPANY_DETAILS['hr_title']}")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "05_New_Employee_Welcome_Letter.docx"))
    
    # === 06_New_Hire_Onboarding_Checklist.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("New Hire Onboarding Checklist / Senarai Semak Orientasi Pekerja Baharu", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_paragraph("Nama Pekerja / Employee Name: <<Nama Pekerja>>")
    doc.add_paragraph("Jawatan / Position: <<Jawatan>>")
    doc.add_paragraph("Tarikh Mula / Start Date: <<Tarikh Mula>>")
    doc.add_paragraph("Pengurus / Manager: <<Nama Pengurus>>")
    doc.add_paragraph()
    doc.add_heading("Fasa 1: Pra-Kemasukan (Sebelum Hari Pertama) / Phase 1: Pre-Boarding (Before Day 1)", 1)
    table_p1 = doc.add_table(rows=1, cols=2); table_p1.style = 'Table Grid'
    add_table_row(table_p1, ["Status", "Tugasan / Task (Ditugaskan kepada / Assigned to: HR/Pengurus)"])
    add_table_row(table_p1, ["[]", "EN: Send official Letter of Offer. Receive signed copy.\nBM: Hantar Surat Tawaran rasmi. Terima salinan yang ditandatangani."])
    add_table_row(table_p1, ["[]", "EN: Collect new hire documents (IC/Passport copy, bank account details...)\nBM: Kumpul dokumen pekerja baharu (Salinan K/P/Pasport, butiran akaun bank...)."])
    add_table_row(table_p1, ["[]", "EN: Register employee for KWSP (EPF) (within 7 days...)\nBM: Daftar pekerja untuk KWSP (EPF) (dalam tempoh 7 hari)."])
    add_table_row(table_p1, ["[]", "EN: Register employee for PERKESO (SOCSO) & EIS (within 30 days...)\nBM: Daftar pekerja untuk PERKESO (SOCSO) & EIS (dalam tempoh 30 hari)."])
    add_table_row(table_p1, ["[]", "EN: Prepare workstation...\nBM: Sediakan stesen kerja..."])
    add_table_row(table_p1, ["[]", "EN: Prepare IT assets... and create accounts...\nBM: Sediakan aset IT... dan cipta akaun..."])
    add_table_row(table_p1, ["[]", "EN: Send \"Welcome Letter\" with first-day details...\nBM: Hantar \"Surat Aluan\" dengan butiran hari pertama..."])
    add_table_row(table_p1, ["[]", "EN: Announce new hire... to the team/company.\nBM: Umumkan pekerja baharu... kepada pasukan/syarikat."])
    doc.add_paragraph()
    doc.add_heading("Fasa 2: Hari Pertama / Phase 2: First Day", 1)
    table_p2 = doc.add_table(rows=1, cols=2); table_p2.style = 'Table Grid'
    add_table_row(table_p2, ["Status", "Tugasan / Task (Ditugaskan kepada / Assigned to: Pengurus/HR)"])
    add_table_row(table_p2, ["[]", "EN: Greet new hire personally upon arrival.\nBM: Sambut pekerja baharu secara peribadi..."])
    add_table_row(table_p2, ["[]", "EN: HR Orientation: Provide Employee Handbook, explain key policies...\nBM: Orientasi HR: Berikan Buku Panduan Pekerja, terangkan polisi utama..."])
    add_table_row(table_p2, ["[]", "EN: Get all remaining HR forms signed...\nBM: Dapatkan tandatangan untuk semua borang HR..."])
    add_table_row(table_p2, ["[]", "EN: Office tour: Introduce workstation, pantry, restrooms, prayer room (surau).\nBM: Lawatan pejabat: Tunjukkan stesen kerja, pantri, tandas, surau."])
    add_table_row(table_p2, ["[]", "EN: Team introductions. Assign an \"onboarding buddy\"...\nBM: Sesi perkenalan dengan pasukan. Lantik \"rakan onboarding\"..."])
    add_table_row(table_p2, ["[]", "EN: IT Setup: Ensure logins are working...\nBM: Persediaan IT: Pastikan log masuk berfungsi..."])
    add_table_row(table_p2, ["[]", "EN: Manager 1-on-1: Discuss first-week goals...\nBM: Sesi 1-dengan-1 Pengurus: Bincang matlamat minggu pertama..."])
    add_table_row(table_p2, ["[]", "EN: Arrange team lunch...\nBM: Aturkan makan tengah hari bersama pasukan..."])
    doc.add_paragraph()
    doc.add_heading("Fasa 3: Bulan Pertama (Hari 2 - 30) / Phase 3: First Month (Day 2 - 30)", 1)
    table_p3 = doc.add_table(rows=1, cols=2); table_p3.style = 'Table Grid'
    add_table_row(table_p3, ["Status", "Tugasan / Task (Ditugaskan kepada / Assigned to: Pengurus)"])
    add_table_row(table_p3, ["[]", "EN: Schedule regular... check-ins with the new hire.\nBM: Jadualkan sesi 'check-in' yang kerap..."])
    add_table_row(table_p3, ["[]", "EN: Set clear 30-day goals and discuss first project.\nBM: Tetapkan matlamat 30-hari yang jelas..."])
    add_table_row(table_p3, ["[]", "EN: Provide necessary job-specific training...\nBM: Berikan latihan khusus untuk kerja..."])
    add_table_row(table_p3, ["[]", "EN: Schedule introductory meetings with key colleagues...\nBM: Jadualkan mesyuarat pengenalan dengan rakan sekerja utama..."])
    add_table_row(table_p3, ["[]", "EN: Conduct an informal \"End of First Week\" check-in...\nBM: Adakan 'check-in' tidak rasmi \"Penghujung Minggu Pertama\"..."])
    add_table_row(table_p3, ["[]", "EN: Conduct the first formal Probationary Review (e.g., at Day 30).\nBM: Lakukan Semakan Percubaan rasmi yang pertama..."])
    add_table_row(table_p3, ["[]", "EN: Ask for feedback on the onboarding process.\nBM: Minta maklum balas tentang proses onboarding."])
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "06_New_Hire_Onboarding_Checklist.docx"))
    
    print("...Hiring Kit DONE.")
    return True # Indicate success

def make_handbook(COMPANY_DETAILS, BRAND_TAGLINE, logo_path, save_folder):
    """Generates the complete Employee Handbook."""
    print("Generating Employee Handbook...")
    doc = Document()
    add_logo(doc, logo_path)
    doc.add_heading("Employee Handbook / Buku Panduan Pekerja", 0)
    disclaimer_en = "This document is a template and not legal advice. This handbook is a guide to the Company's policies and procedures; it is not an employment contract, although some policies reflect statutory or contractual terms. The Company reserves the right to amend policies at any time."
    disclaimer_bm = "Dokumen ini hanyalah templat dan bukan nasihat undang-undang. Buku panduan ini adalah panduan kepada polisi dan prosedur Syarikat; ia bukan kontrak pekerjaan, walaupun sesetengah polisi mencerminkan terma statutori atau kontraktual. Syarikat berhak untuk meminda polisi pada bila-bila masa."
    add_bilingual_disclaimer(doc, disclaimer_en, disclaimer_bm)
    
    doc.add_heading("JADUAL KANDUNGAN / TABLE OF CONTENTS", level=1)
    doc.add_paragraph(
        "1.0 PENGENALAN / INTRODUCTION\n"
        "  1.1 Mesej Aluan / Welcome Message\n"
        "  1.2 Misi & Nilai Syarikat / Company Mission & Values\n"
        "2.0 POLISI PEKERJAAN / EMPLOYMENT POLICIES\n"
        "  2.1 Tempoh Percubaan / Probationary Period\n"
        "  2.2 Pengesahan / Confirmation\n"
        "  2.3 Sulit & Data Peribadi / Confidentiality & Personal Data\n"
        "  2.4 Susunan Kerja Fleksibel (FWA) / Flexible Work Arrangements\n"
        "3.0 WAKTU BEKERJA & KERJA LEBIH MASA / WORKING HOURS & OVERTIME\n"
        "  3.1 Waktu Bekerja / Working Hours\n"
        "  3.2 Kerja Lebih Masa (Overtime) / Overtime\n"
        "4.0 CUTI-CUTI DIPERUNTUKKAN / LEAVE ENTITLEMENTS\n"
        "  4.1 Cuti Tahunan / Annual Leave\n"
        "  4.2 Cuti Sakit / Sick Leave\n"
        "  4.3 Cuti Hospitalisasi / Hospitalisation Leave\n"
        "  4.4 Cuti Bersalin / Maternity Leave\n"
        "  4.5 Cuti Paterniti / Paternity Leave\n"
        "  4.6 Cuti Umum / Public Holidays\n"
        "  4.7 Cuti Ehsan / Compassionate Leave\n"
        "5.0 PAMPASAN & FAEDAH / COMPENSATION & BENEFITS\n"
        "  5.1 Pembayaran Gaji / Salary Payment\n"
        "  5.2 Potongan Berkanun (KWSP, PERKESO, EIS) / Statutory Deductions\n"
        "  5.3 Tuntutan / Claims\n"
        "6.0 TATAKELAKUAN & KESELAMATAN TEMPAT KERJA / WORKPLACE CONDUCT & SAFETY\n"
        "  6.1 Kod Kelakuan Profesional / Code of Professional Conduct\n"
        "  6.2 Kod Pakaian / Dress Code\n"
        "  6.3 Anti-Gangguan Seksual / Anti-Sexual Harassment\n"
        "  6.4 Anti-Buli & Diskriminasi / Anti-Bullying & Discrimination\n"
        "  6.5 Kesihatan & Keselamatan / Health & Safety\n"
        "7.0 PROSEDUR TATATERTIB / DISCIPLINARY PROCEDURES\n"
        "  7.1 Am / General\n"
        "  7.2 Jenis Salah Laku / Types of Misconduct\n"
        "  7.3 Prosedur Tatatertib (Due Inquiry) / Disciplinary Procedure\n"
        "8.0 PENAMATAN PERKHIDMATAN / TERMINATION OF EMPLOYMENT\n"
        "  8.1 Notis Penamatan / Notice of Termination\n"
        "  8.2 Penamatan Serta-Merta (Salah Laku) / Summary Dismissal (Misconduct)\n"
        "9.0 HALAMAN AKUAN PEKERJA / EMPLOYEE ACKNOWLEDGEMENT PAGE"
    )
    doc.add_page_break()

    # --- 1.0 PENGENALAN / INTRODUCTION ---
    doc.add_heading("1.0 PENGENALAN / INTRODUCTION", level=1)
    doc.add_heading("1.1 Mesej Aluan / Welcome Message", level=2)
    add_bilingual_block(doc,
        f"Welcome to {COMPANY_DETAILS['name']}! Whether you are new to our team or have been with us for a while, this handbook serves as your guide to understanding our culture, policies, and the expectations we share. Our success is built on our people, and we are committed to creating a positive and productive environment for everyone.",
        f"Selamat datang ke {COMPANY_DETAILS['name']}! Sama ada anda baharu dalam pasukan kami atau telah lama bersama kami, buku panduan ini berfungsi sebagai panduan anda untuk memahami budaya, polisi, dan jangkaan yang kami kongsi bersama. Kejayaan kami dibina atas bakat warga kerja kami, dan kami komited untuk mewujudkan persekitaran yang positif dan produktif untuk semua."
    )
    doc.add_paragraph(f"\n{COMPANY_DETAILS['ceo_name']}\n{COMPANY_DETAILS['ceo_title']}")

    doc.add_heading("1.2 Misi & Nilai Syarikat / Company Mission & Values", level=2)
    add_bilingual_block(doc,
        "Our Mission: <<Insert Company Mission>>\nOur Values:\n• <<Value 1 (e.g., Integrity)>>\n• <<Value 2 (e.g., Customer First)>>\n• <<Value 3 (e.g., Teamwork)>>",
        "Misi Kami: <<Insert Misi Syarikat>>\nNilai-Nilai Kami:\n• <<Nilai 1 (cth., Integriti)>>\n• <<Nilai 2 (cth., Pelanggan Didahulukan)>>\n• <<Nilai 3 (cth., Kerja Berpasukan)>>"
    )

    # --- 2.0 POLISI PEKERJAAN / EMPLOYMENT POLICIES ---
    doc.add_heading("2.0 POLISI PEKERJAAN / EMPLOYMENT POLICIES", level=1)
    doc.add_heading("2.1 Tempoh Percubaan / Probationary Period", level=2)
    add_bilingual_block(doc,
        "All new employees will serve a probationary period of <<Three (3) to Six (6)>> months. This period is for the Company to assess your suitability for the role and for you to evaluate the Company. This period may be extended at the Company's sole discretion.",
        "Semua pekerja baharu akan menjalani tempoh percubaan selama <<Tiga (3) hingga Enam (6)>> bulan. Tempoh ini adalah untuk Syarikat menilai kesesuaian anda untuk jawatan tersebut dan untuk anda menilai Syarikat. Tempoh ini boleh dilanjutkan atas budi bicara Syarikat semata-mata."
    )
    doc.add_heading("2.2 Pengesahan / Confirmation", level=2)
    add_bilingual_block(doc, "Upon successful completion of the probationary period, your employment will be confirmed in writing.", "Selepas berjaya menamatkan tempoh percubaan, perkhidmatan anda akan disahkan secara bertulis.")

    doc.add_heading("2.3 Sulit & Data Peribadi / Confidentiality & Personal Data", level=2)
    add_bilingual_block(doc,
        "During your employment, you will have access to confidential information. You must not disclose this information to any third party, during or after your employment. The Company respects your personal data in accordance with the Personal Data Protection Act 2010 (PDPA).",
        "Sepanjang perkhidmatan anda, anda akan mempunyai akses kepada maklumat sulit. Anda tidak boleh mendedahkan maklumat ini kepada mana-mana pihak ketiga, semasa atau selepas perkhidmatan anda. Syarikat menghormati data peribadi anda selaras dengan Akta Perlindungan Data Peribadi 2010 (PDPA)."
    )
    doc.add_heading("2.4 Susunan Kerja Fleksibel (FWA) / Flexible Work Arrangements", level=2)
    add_bilingual_block(doc,
        "In line with the Employment (Amendment) Act 2022, all employees may apply for a Flexible Work Arrangement (FWA), such as variations in working hours, days, or location. Applications must be made in writing to your Head of Department. The Company will provide a written decision within 60 days of the application. Approval is subject to business and operational needs.",
        "Selaras dengan Akta Kerja (Pindaan) 2022, semua pekerja boleh memohon Susunan Kerja Fleksibel (FWA), seperti perubahan dalam waktu kerja, hari bekerja, atau lokasi kerja. Permohonan mesti dibuat secara bertulis kepada Ketua Jabatan anda. Syarikat akan memberikan keputusan bertulis dalam tempoh 60 hari dari tarikh permohonan. Kelulusan adalah tertakluk kepada keperluan perniagaan dan operasi."
    )

    # --- 3.0 WAKTU BEKERJA & KERJA LEBIH MASA ---
    doc.add_heading("3.0 WAKTU BEKERJA & KERJA LEBIH MASA / WORKING HOURS & OVERTIME", level=1)
    doc.add_heading("3.1 Waktu Bekerja / Working Hours", level=2)
    add_bilingual_block(doc,
        f"The official working hours for the Company are from {COMPANY_DETAILS['working_hours']}. This constitutes a total of 45 hours per week, in compliance with the Employment Act 1955.",
        f"Waktu bekerja rasmi Syarikat adalah dari {COMPANY_DETAILS['working_hours']}. Ini bersamaan dengan jumlah 45 jam seminggu, mematuhi Akta Kerja 1955."
    )
    doc.add_heading("3.2 Kerja Lebih Masa (Overtime) / Overtime", level=2)
    add_bilingual_block(doc,
        "1. Employees earning RM4,000 or less per month (and other categories specified in the Act): You are eligible for overtime (OT) payments as per the rates prescribed in the Employment Act 1955. All OT must be approved in advance by your Manager.",
        "1. Pekerja bergaji RM4,000 atau kurang sebulan (dan kategori lain yang dinyatakan dalam Akta): Anda layak untuk bayaran kerja lebih masa (OT) mengikut kadar yang ditetapkan dalam Akta Kerja 1955. Semua OT mesti diluluskan terlebih dahulu oleh Pengurus anda."
    )
    add_bilingual_block(doc,
        "2. Employees earning more than RM4,000 per month: You are not eligible for OT payments. However, based on operational needs and at your Manager's discretion, 'time-off in lieu' (replacement leave) may be granted for excessive hours worked.",
        "2. Pekerja bergaji lebih daripada RM4,000 sebulan: Anda tidak layak untuk bayaran OT. Walau bagaimanapun, berdasarkan keperluan operasi dan atas budi bicara Pengurus, 'cuti ganti' (replacement leave) mungkin diberikan untuk waktu kerja berlebihan."
    )

    # --- 4.0 CUTI-CUTI DIPERUNTUKKAN / LEAVE ENTITLEMENTS ---
    doc.add_heading("4.0 CUTI-CUTI DIPERUNTUKKAN / LEAVE ENTITLEMENTS", level=1)
    doc.add_heading("4.1 Cuti Tahunan / Annual Leave", level=2)
    add_bilingual_block(doc,
        "Paid annual leave entitlement is based on your length of service:\n• Less than 2 years: 8 days\n• 2 to 5 years: 12 days\n• More than 5 years: 16 days",
        "Kelayakan cuti tahunan berbayar adalah berdasarkan tempoh perkhidmatan anda:\n• Kurang dari 2 tahun: 8 hari\n• 2 hingga 5 tahun: 12 hari\n• Lebih dari 5 tahun: 16 hari"
    )
    doc.add_heading("4.2 Cuti Sakit / Sick Leave", level=2)
    add_bilingual_block(doc,
        "Paid sick leave (non-hospitalisation) is as follows, provided you notify your Manager and submit a valid Medical Certificate (MC):\n• Less than 2 years: 14 days\n• 2 to 5 years: 18 days\n• More than 5 years: 22 days",
        "Cuti sakit berbayar (bukan hospitalisasi) adalah seperti berikut, dengan syarat anda memaklumkan Pengurus dan mengemukakan Sijil Cuti Sakit (MC) yang sah:\n• Kurang dari 2 tahun: 14 hari\n• 2 hingga 5 tahun: 18 hari\n• Lebih dari 5 tahun: 22 hari"
    )
    doc.add_heading("4.3 Cuti Hospitalisasi / Hospitalisation Leave", level=2)
    add_bilingual_block(doc,
        "You are entitled to 60 days of paid hospitalisation leave per calendar year. This is a separate entitlement from the sick leave mentioned in 4.2.",
        "Anda layak mendapat 60 hari cuti hospitalisasi berbayar setiap tahun kalendar. Ini adalah kelayakan yang berasingan daripada cuti sakit yang dinyatakan dalam 4.2."
    )
    doc.add_heading("4.4 Cuti Bersalin / Maternity Leave", level=2)
    add_bilingual_block(doc,
        "Eligible female employees are entitled to 98 consecutive days of paid maternity leave.",
        "Pekerja wanita yang layak berhak mendapat cuti bersalin berbayar selama 98 hari berturut-turut."
    )
    doc.add_heading("4.5 Cuti Paterniti / Paternity Leave", level=2)
    add_bilingual_block(doc,
        "Eligible married male employees are entitled to 7 consecutive days of paid paternity leave per confinement, limited to five (5) confinements. Eligibility requires at least 12 months of service.",
        "Pekerja lelaki yang telah berkahwin yang layak berhak mendapat 7 hari berturut-turut cuti paterniti berbayar untuk setiap kelahiran, terhad kepada lima (5) kelahiran. Kelayakan memerlukan sekurang-kurangnya 12 bulan perkhidmatan."
    )
    doc.add_heading("4.6 Cuti Umum / Public Holidays", level=2)
    add_bilingual_block(doc,
        "The Company observes 11 gazetted public holidays per year, as mandated by the Act. 5 of these are compulsory (National Day, YDPA's Birthday, State Ruler's Birthday/FT Day, Labour Day, Malaysia Day). The remaining 6 will be announced by HR.",
        "Syarikat akan mematuhi 11 hari cuti umum yang diwartakan setiap tahun, seperti yang dimandatkan oleh Akta. 5 daripadanya adalah wajib (Hari Kebangsaan, Hari Keputeraan YDPA, Hari Keputeraan Raja Negeri/Hari Wilayah Persekutuan, Hari Pekerja, Hari Malaysia). Baki 6 hari lagi akan diumumkan oleh HR."
    )
    doc.add_heading("4.7 Cuti Ehsan / Compassionate Leave", level=2)
    add_bilingual_block(doc,
        "The Company provides paid compassionate (bereavement) leave of <<e.g., 3>> days in the event of the passing of an immediate family member (spouse, child, parent, or sibling).",
        "Syarikat menyediakan cuti ehsan berbayar (kematian) selama <<cth., 3>> hari sekiranya berlaku kematian ahli keluarga terdekat (pasangan, anak, ibu bapa, atau adik-beradik)."
    )

    # --- 5.0 PAMPASAN & FAEDAH / COMPENSATION & BENEFITS ---
    doc.add_heading("5.0 PAMPASAN & FAEDAH / COMPENSATION & BENEFITS", level=1)
    doc.add_heading("5.1 Pembayaran Gaji / Salary Payment", level=2)
    add_bilingual_block(doc, "Your salary will be paid on a monthly basis, no later than the 7th day of the following month.", "Gaji anda akan dibayar secara bulanan, tidak lewat daripada 7 haribulan berikutnya.")
    doc.add_heading("5.2 Potongan Berkanun (KWSP, PERKESO, EIS) / Statutory Deductions", level=2)
    add_bilingual_block(doc, "The Company will make all statutory contributions and deductions (EPF, SOCSO, EIS) as required by law.", "Syarikat akan membuat semua caruman dan potongan berkanun (KWSP, PERKESO, EIS) seperti yang dikehendaki oleh undang-undang.")
    doc.add_heading("5.3 Tuntutan / Claims", level=2)
    add_bilingual_block(doc, "Employees are entitled to claim for work-related expenses (e.g., travel, tolls) as per the Company's claims policy. Please refer to HR for the full policy.", "Pekerja layak menuntut perbelanjaan berkaitan kerja (cth., perjalanan, tol) mengikut polisi tuntutan Syarikat. Sila rujuk HR untuk polisi penuh.")
    
    # --- 6.0 TATAKELAKUAN & KESELAMATAN TEMPAT KERJA ---
    doc.add_heading("6.0 TATAKELAKUAN & KESELAMATAN TEMPAT KERJA / WORKPLACE CONDUCT & SAFETY", level=1)
    doc.add_heading("6.1 Kod Kelakuan Profesional / Code of Professional Conduct", level=2)
    add_bilingual_block(doc,
        "All employees are expected to act with integrity, professionalism, and respect towards colleagues, clients, and suppliers. This includes protecting company property and confidential information.",
        "Semua pekerja dijangka bertindak dengan integriti, profesionalisme, dan rasa hormat terhadap rakan sekerja, pelanggan, dan pembekal. Ini termasuk melindungi harta syarikat dan maklumat sulit."
    )
    doc.add_heading("6.2 Kod Pakaian / Dress Code", level=2)
    add_bilingual_block(doc,
        f"Employees must maintain a neat, professional, and clean appearance (\"{COMPANY_DETAILS['dress_code']}\"). Grooming styles dictated by religion and ethnicity are permitted, provided they are neat and do not pose a safety hazard. T-shirts, shorts, and flip-flops are not permitted.",
        f"Pekerja mesti mengekalkan penampilan yang kemas, profesional, dan bersih (\"{COMPANY_DETAILS['dress_code']}\"). Gaya dandanan atas dasar agama dan etnik adalah dibenarkan, asalkan ia kemas dan tidak menimbulkan bahaya keselamatan. Baju-T, seluar pendek, dan selipar adalah tidak dibenarkan."
    )
    doc.add_heading("6.3 Anti-Gangguan Seksual / Anti-Sexual Harassment", level=2)
    add_bilingual_block(doc,
        "The Company has a zero-tolerance policy for sexual harassment in any form (verbal, physical, visual, or otherwise). This is a serious offence. Any employee who feels harassed must report it immediately to HR or Management. All reports will be investigated promptly and confidentially.",
        "Syarikat mempunyai polisi toleransi sifar terhadap gangguan seksual dalam apa jua bentuk (lisan, fizikal, visual, atau lain-lain). Ini adalah kesalahan yang serius. Mana-mana pekerja yang berasa diganggu mesti melaporkannya dengan segera kepada HR atau Pengurusan. Semua laporan akan disiasat dengan segera dan sulit."
    )
    doc.add_heading("6.4 Anti-Buli & Diskriminasi / Anti-Bullying & Discrimination", level=2)
    add_bilingual_block(doc, "The Company is committed to a workplace free of bullying, discrimination, and harassment. All employees must be treated with dignity and respect.", "Syarikat komited kepada tempat kerja yang bebas daripada buli, diskriminasi, dan gangguan. Semua pekerja mesti dilayan dengan maruah dan rasa hormat.")
    doc.add_heading("6.5 Kesihatan & Keselamatan / Health & Safety", level=2)
    add_bilingual_block(doc, "The Company is committed to providing a safe and healthy work environment. Employees must comply with all safety rules and report any unsafe conditions to Management.", "Syarikat komited untuk menyediakan persekitaran kerja yang selamat dan sihat. Pekerja mesti mematuhi semua peraturan keselamatan dan melaporkan sebarang keadaan yang tidak selamat kepada Pengurusan.")

    # --- 7.0 PROSEDUR TATATERTIB / DISCIPLINARY PROCEDURES ---
    doc.add_heading("7.0 PROSEDUR TATATERTIB / DISCIPLINARY PROCEDURES", level=1)
    doc.add_heading("7.1 Am / General", level=2)
    add_bilingual_block(doc,
        "The Company's goal is to correct behaviour, not to punish. However, to ensure fairness and safety, disciplinary action is necessary for misconduct. All actions will be based on \"just cause and excuse\" as required by the Industrial Relations Act 1967.",
        "Matlamat Syarikat adalah untuk membetulkan tingkah laku, bukan untuk menghukum. Walau bagaimanapun, untuk memastikan keadilan dan keselamatan, tindakan tatatertib adalah perlu untuk salah laku. Semua tindakan akan berdasarkan \"alasan yang adil\" seperti yang dikehendaki oleh Akta Perhubungan Perusahaan 1967."
    )
    doc.add_heading("7.2 Jenis Salah Laku / Types of Misconduct", level=2)
    add_bilingual_block(doc,
        "• Minor Misconduct: (e.g., Tardiness, improper attire, non-compliance with simple rules).\n• Major/Gross Misconduct: (e.g., Theft, fraud, fighting, insubordination, sexual harassment, absence from work for more than 2 consecutive days without notice).",
        "• Salah Laku Kecil: (cth., Lewat, pakaian tidak sesuai, kegagalan mematuhi peraturan mudah).\n• Salah Laku Berat/Besar: (cth., Kecurian, penipuan, bergaduh, ingkar arahan, gangguan seksual, tidak hadir bekerja lebih dari 2 hari berturut-turut tanpa notis)."
    )
    doc.add_heading("7.3 Prosedur Tatatertib (Due Inquiry) / Disciplinary Procedure (Due Inquiry)", level=2)
    add_bilingual_block(doc,
        "For major misconduct, the Company will follow a fair \"due inquiry\" process:\n1. Show Cause Letter: The employee will be issued a \"Surat Tunjuk Sebab\" (Show Cause Letter) detailing the allegations and given a reasonable time (e.g., 2-7 days) to provide a written explanation.\n2. Investigation / Inquiry: If the explanation is not satisfactory, the Company may hold a \"Domestic Inquiry\" (DI) (Siasatan Dalaman) to hear the evidence and allow the employee to state their case.\n3. Decision: Based on the findings, the Company will decide on the appropriate punishment, which may include a warning, suspension, or dismissal.",
        "Untuk salah laku berat, Syarikat akan mengikut proses \"due inquiry\" (siasatan wajar) yang adil:\n1. Surat Tunjuk Sebab: Pekerja akan diberikan \"Surat Tunjuk Sebab\" yang memperincikan pertuduhan dan diberi masa yang munasabah (cth., 2-7 hari) untuk memberi penjelasan bertulis.\n2. Siasatan / Inkuiri: Jika penjelasan tidak memuaskan, Syarikat boleh mengadakan \"Siasatan Dalaman\" (Domestic Inquiry - DI) untuk mendengar bukti dan membenarkan pekerja membela diri.\n3. Keputusan: Berdasarkan penemuan siasatan, Syarikat akan memutuskan hukuman yang sewajarnya, yang mungkin termasuk amaran, penggantungan, atau pemecatan."
    )

    # --- 8.0 PENAMATAN PERKHIDMATAN / TERMINATION OF EMPLOYMENT ---
    doc.add_heading("8.0 PENAMATAN PERKHIDMATAN / TERMINATION OF EMPLOYMENT", level=1)
    doc.add_heading("8.1 Notis Penamatan / Notice of Termination", level=2)
    add_bilingual_block(doc,
        "After confirmation, the notice period required by either party to terminate employment is as follows:\n• Less than 2 years of service: 4 weeks\n• 2 to 5 years of service: 6 weeks\n• More than 5 years of service: 8 weeks",
        "Selepas pengesahan, tempoh notis yang diperlukan oleh mana-mana pihak untuk menamatkan perkhidmatan adalah seperti berikut:\n• Kurang dari 2 tahun perkhidmatan: 4 minggu\n• 2 hingga 5 tahun perkhidmatan: 6 minggu\n• Lebih dari 5 tahun perkhidmatan: 8 minggu"
    )
    doc.add_heading("8.2 Penamatan Serta-Merta (Salah Laku) / Summary Dismissal (Misconduct)", level=2)
    add_bilingual_block(doc,
        "The Company reserves the right to terminate your employment without notice (summary dismissal) in cases of proven gross misconduct, after a due inquiry process has been conducted.",
        "Syarikat berhak menamatkan perkhidmatan anda tanpa notis (pemecatan serta-merta) dalam kes salah laku berat yang telah dibuktikan, selepas proses siasatan wajar (due inquiry) dijalankan."
    )

    # --- 9.0 HALAMAN AKUAN PEKERJA / EMPLOYEE ACKNOWLEDGEMENT PAGE ---
    doc.add_heading("9.0 HALAMAN AKUAN PEKERJA / EMPLOYEE ACKNOWLEDGEMENT PAGE", level=1)
    doc.add_page_break()
    
    add_bilingual_block(doc,
        f"I, <<Employee's Full Name>> (NRIC No: <<NRIC No.>>), hereby acknowledge that I have received, read, and understood the policies and procedures outlined in the {COMPANY_DETAILS['name']} Employee Handbook (Version <<1.0>>). I agree to abide by these terms and conditions as a condition of my employment. I understand that this Handbook is a guide and that the Company reserves the right to amend its contents at any time.",
        f"Saya, <<Nama Penuh Pekerja>> (No. K/P: <<No. K/P>>), dengan ini mengaku bahawa saya telah menerima, membaca, dan memahami polisi serta prosedur yang terkandung di dalam Buku Panduan Pekerja {COMPANY_DETAILS['name']} (Version <<1.0>>). Saya bersetuju untuk mematuhi terma dan syarat ini sebagai syarat pekerjaan saya. Saya faham bahawa Buku Panduan ini adalah panduan dan Syarikat berhak untuk meminda kandungannya pada bila-bila masa."
    )
    doc.add_paragraph("\n\n________________________________________")
    doc.add_paragraph("Tandatangan Pekerja / Employee Signature:")
    doc.add_paragraph("\nNama / Name: ___________________________")
    doc.add_paragraph("\nTarikh / Date: ___________________________")

    # --- Footer and Save ---
    add_footer(doc, BRAND_TAGLINE)
    doc.save(os.path.join(save_folder, "01_Employee_Handbook_Template.docx"))
    print("...Employee Handbook DONE.")
    return True # Indicate success

def make_performance(COMPANY_DETAILS, BRAND_TAGLINE, logo_path, save_folder):
    """Generates all 5 documents for the Performance Management Toolkit."""
    print("Generating Performance Toolkit...")
    
    # === 01_Performance_Review_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Performance Review Template / Templat Penilaian Prestasi", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_heading("BAHAGIAN A: MAKLUMAT PEKERJA / SECTION A: EMPLOYEE DETAILS", 1)
    table_a = doc.add_table(rows=4, cols=2); table_a.style = 'Table Grid'
    table_a.cell(0, 0).text = "Nama Pekerja / Employee Name:"; table_a.cell(1, 0).text = "Jawatan / Position:"
    table_a.cell(2, 0).text = "Pengurus / Manager:"; table_a.cell(3, 0).text = "Tempoh Penilaian / Review Period:"
    table_a.cell(3, 1).text = "Dari / From: <<...>> Hingga / To: <<...>>"
    doc.add_paragraph()
    doc.add_heading("BAHAGIAN B: PENILAIAN MATLAMAT / SECTION B: GOAL ASSESSMENT", 1)
    table_b = doc.add_table(rows=1, cols=3); table_b.style = 'Table Grid'
    add_table_row(table_b, ["Matlamat / Goal", "Hasil (Pencapaian) / Result (Achievement)", "Ulasan Pengurus / Manager's Comments"])
    add_table_row(table_b, ["1. <<Goal 1>>", "", ""]); add_table_row(table_b, ["2. <<Goal 2>>", "", ""]); add_table_row(table_b, ["3. <<Goal 3>>", "", ""])
    doc.add_paragraph()
    doc.add_heading("BAHAGIAN C: PENILAIAN KOMPETENSI / SECTION C: COMPETENCY ASSESSMENT", 1)
    doc.add_paragraph("Skala Penilaian / Rating Scale:\n1 = Perlu Penambahbaikan / Needs Improvement\n2 = Memenuhi Jangkaan / Meets Expectations\n3 = Melebihi Jangkaan / Exceeds Expectations")
    table_c = doc.add_table(rows=1, cols=3); table_c.style = 'Table Grid'
    add_table_row(table_c, ["Kompetensi / Competency", "Penilaian (1-3) / Rating (1-3)", "Ulasan & Contoh / Comments & Examples"])
    add_table_row(table_c, ["Kualiti Kerja / Quality of Work\n(Ketepatan, teliti / Accuracy, thoroughness)", "", ""])
    add_table_row(table_c, ["Produktiviti / Productivity\n(Pengurusan masa, kuantiti kerja / Time management, output)", "", ""])
    add_table_row(table_c, ["Komunikasi / Communication\n(Jelas, responsif / Clarity, responsiveness)", "", ""])
    add_table_row(table_c, ["Kerja Berpasukan / Teamwork\n(Bekerjasama, menyokong / Collaborative, supportive)", "", ""])
    add_table_row(table_c, ["Inisiatif / Initiative\n(Penyelesaian masalah, proaktif / Problem-solving, proactive)", "", ""])
    doc.add_paragraph()
    doc.add_heading("BAHAGIAN D: ULASAN & PELAN PEMBANGUNAN / SECTION D: COMMENTS & DEVELOPMENT PLAN", 1)
    doc.add_paragraph("1. Ulasan Keseluruhan Pengurus / Manager's Overall Comments:\n(Kekuatan utama & bidang utama untuk penambahbaikan)\n<<...>>\n")
    doc.add_paragraph("2. Ulasan Pekerja / Employee's Comments:\n(Komen mengenai penilaian ini)\n<<...>>\n")
    doc.add_paragraph("3. Pelan Pembangunan / Development Plan:\n(Matlamat & latihan untuk tempoh seterusnya)\n<<...>>\n")
    doc.add_heading("BAHAGIAN E: AKUAN / SECTION E: ACKNOWLEDGEMENT", 1)
    add_bilingual_block(doc, "We have discussed this review and I have received a copy. My signature does not necessarily imply agreement.", "Kami telah membincangkan penilaian ini dan saya telah menerima satu salinan. Tandatangan saya tidak semestinya menandakan persetujuan.")
    doc.add_paragraph("\n\n___________________\nTandatangan Pekerja / Employee Signature:\nTarikh / Date:\n\n")
    doc.add_paragraph("___________________\nTandatangan Pengurus / Manager Signature:\nTarikh / Date:")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "01_Performance_Review_Template.docx"))

    # === 02_Employee_Self_Evaluation_Form.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Employee Self-Evaluation Form / Borang Penilaian Kendiri Pekerja", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_paragraph()
    table_se = doc.add_table(rows=3, cols=2)
    table_se.cell(0, 0).text = "Nama Pekerja / Employee Name:"; table_se.cell(1, 0).text = "Jawatan / Position:"; table_se.cell(2, 0).text = "Tempoh Penilaian / Review Period:"
    doc.add_paragraph()
    doc.add_heading("1. Pencapaian Terbesar Saya / My Biggest Achievements", 1)
    doc.add_paragraph("(Senaraikan 3-5 pencapaian utama anda dalam tempoh ini.)\n<<...>>\n")
    doc.add_heading("2. Cabaran yang Saya Hadapi / Challenges I Faced", 1)
    doc.add_paragraph("(Apakah halangan yang anda temui?)\n<<...>>\n")
    doc.add_heading("3. Bidang yang Saya Ingin Perbaiki / Areas I Want to Improve", 1)
    doc.add_paragraph("(Apakah kemahiran atau pengetahuan yang ingin anda bangunkan?)\n<<...>>\n")
    doc.add_heading("4. Sokongan yang Saya Perlukan / Support I Need from My Manager", 1)
    doc.add_paragraph("(Bagaimana pengurus anda boleh bantu anda untuk berjaya? cth., latihan, maklum balas.)\n<<...>>\n")
    doc.add_heading("5. Matlamat Cadangan Saya untuk Tempoh Seterusnya / My Proposed Goals for the Next Period", 1)
    doc.add_paragraph("<<...>>\n")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "02_Employee_Self_Evaluation_Form.docx"))

    # === 03_SMART_Goals_OKR_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("SMART Goals & OKR Template", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    doc.add_heading("BAHAGIAN 1: MATLAMAT S.M.A.R.T. / SECTION 1: S.M.A.R.T. GOALS", 1)
    add_bilingual_block(doc, "SMART is a framework for setting clear, individual goals...", "SMART ialah rangka kerja untuk menetapkan matlamat individu yang jelas...")
    doc.add_paragraph("• S - Specific / Spesifik\n• M - Measurable / Boleh Diukur\n• A - Achievable / Boleh Dicapai\n• R - Relevant / Relevan\n• T - Time-bound / Tempoh Masa")
    doc.add_heading("Templat Matlamat SMART / SMART Goal Template:", 2)
    doc.add_paragraph("Matlamat / Goal: <<Ringkasan Matlamat>>\nS (Spesifik): <<...>>\nM (Boleh Diukur): <<...>>\nA (Boleh Dicapai): <<...>>\nR (Relevan): <<...>>\nT (Tempoh Masa): <<...>>")
    doc.add_paragraph()
    doc.add_heading("BAHAGIAN 2: OBJEKTIF & HASIL UTAMA (OKR) / SECTION 2: OBJECTIVES & KEY RESULTS (OKR)", 1)
    add_bilingual_block(doc, "OKR is a framework for setting ambitious, collaborative goals...", "OKR ialah rangka kerja untuk menetapkan matlamat kolaboratif...")
    doc.add_heading("Templat OKR / OKR Template:", 2)
    doc.add_paragraph("Objektif / Objective:\n(Cth: Tingkatkan kesedaran jenama PKS kami secara online)\n<<...>>\n")
    doc.add_paragraph("Hasil Utama / Key Results:\n(Mesti boleh diukur)\n1. HR 1: <<Contoh: Capai 10,000 pengikut Instagram...>>\n2. HR 2: <<Contoh: Terbitkan 4 catatan blog...>>\n3. HR 3: <<Contoh: Dapatkan 5 liputan media...>>")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "03_SMART_Goals_OKR_Template.docx"))

    # === 04_Manager_Employee_1-on-1_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Manager-Employee 1-on-1 Template / Templat Mesyuarat 1-dengan-1", 0)
    add_bilingual_disclaimer(doc, "This document is a template and not legal advice.", "Dokumen ini hanyalah templat dan bukan nasihat undang-undang.")
    add_bilingual_block(doc, "To facilitate regular, informal check-ins... This is for coaching and support, not evaluation.", "Untuk memudahcarakan 'check-in' tidak rasmi... Ini adalah untuk bimbingan dan sokongan, bukan penilaian.")
    doc.add_paragraph()
    table_1on1 = doc.add_table(rows=3, cols=2)
    table_1on1.cell(0, 0).text = "Pekerja / Employee:"; table_1on1.cell(1, 0).text = "Pengurus / Manager:"; table_1on1.cell(2, 0).text = "Tarikh / Date:"
    doc.add_paragraph()
    doc.add_heading("1. Agenda / Perkara Perbincangan", 1); doc.add_paragraph("<<...>>\n")
    doc.add_heading("2. Kemajuan & Pencapaian", 1); doc.add_paragraph("<<...>>\n")
    doc.add_heading("3. Cabaran & Halangan", 1); doc.add_paragraph("<<...>>\n")
    doc.add_heading("4. Perkara Tindakan / Action Items", 1); doc.add_paragraph("[ ] <<Tindakan 1>> - Oleh / By: <> Tarikh Siap / Due: <>\n[ ] <<Tindakan 2>> - Oleh / By: <> - Tarikh Siap / Due: <<>")
    doc.add_heading("5. Tarikh Mesyuarat Seterusnya:", 1); doc.add_paragraph("<<...>>")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "04_Manager_Employee_1-on-1_Template.docx"))

    # === 05_Performance_Improvement_Plan_Template.docx ===
    doc = Document(); add_logo(doc, logo_path)
    doc.add_heading("Performance Improvement Plan (PIP) Template", 0)
    add_bilingual_disclaimer(doc, "A PIP must be conducted in \"good faith\"... Failure to do so may lead to claims of unfair dismissal.", "PIP mesti dijalankan dengan \"niat baik\"... Kegagalan berbuat demikian boleh membawa kepada tuntutan pemecatan yang tidak adil.")
    doc.add_paragraph()
    table_pip = doc.add_table(rows=4, cols=2); table_pip.style = 'Table Grid'
    table_pip.cell(0, 0).text = "Nama Pekerja / Employee Name:"; table_pip.cell(1, 0).text = "Jawatan / Position:"
    table_pip.cell(2, 0).text = "Pengurus / Manager:"; table_pip.cell(3, 0).text = "Tarikh Mula PIP / PIP Start Date:"
    table_pip.cell(3, 1).text = "Tarikh Semakan Akhir / Final Review Date:\n<<cth., 60 hari dari tarikh mula / e.g., 60 days from start date>>"
    doc.add_paragraph()
    doc.add_heading("BAHAGIAN 1: BIDANG PENAMBAHBAIKAN KHUSUS / SECTION 1: SPECIFIC AREAS OF IMPROVEMENT", 1)
    add_bilingual_block(doc, "This section must detail the specific performance gaps, with documented examples... NOT \"Is always late\").", "Bahagian ini mesti memperincikan jurang prestasi yang spesifik, dengan contoh... BUKAN \"Selalu lewat\").")
    doc.add_paragraph("1. <<Kelemahan Prestasi 1 / Performance Gap 1>>\n2. <<Kelemahan Prestasi 2 / Performance Gap 2>>\n")
    doc.add_heading("BAHAGIAN 2: OBJEKTIF & JANGKAAN BOLEH DIUKUR / SECTION 2: MEASURABLE OBJECTIVES & EXPECTATIONS", 1)
    doc.add_paragraph("EN: By <<Final Review Date>>, <<Employee Name>> is expected to:\nBM: Menjelang <<Tarikh Semakan Akhir>>, <<Nama Pekerja>> dijangka untuk:")
    doc.add_paragraph("1. <<Objektif 1 (cth., Submit 100% of weekly reports on time...)>>\n2. <<Objektif 2 (cth., Reduce customer complaint emails...)>>\n")
    doc.add_heading("BAHAGIAN 3: SOKONGAN & SUMBER DISEDIAKAN / SECTION 3: SUPPORT & RESOURCES PROVIDED", 1)
    add_bilingual_block(doc, "This is a crucial part of a \"good faith\" PIP...", "Ini adalah bahagian penting dalam PIP \"niat baik\"...")
    doc.add_paragraph("1. <<Sokongan 1 (cth., Weekly 30-minute check-in meetings...)>>\n2. <<Sokongan 2 (cth., Access to Company's online 'Customer Service' training...)>>\n3. <<Sokongan 3 (cth., Mentorship from a senior team member...)>>\n")
    doc.add_heading("BAHAGIAN 4: JADUAL SEMAKAN / SECTION 4: REVIEW SCHEDULE", 1)
    doc.add_paragraph("• Semakan 1/ Check-in 1 (cth., Week 2): <<Tarikh>>\n• Semakan 2/ Check-in 2 (cth., Week 4): <<Tarikh>>\n• Semakan Akhir / Final Review (cth., Week 8): <<Tarikh>>\n")
    doc.add_heading("BAHAGIAN 5: AKIBAT JIKA GAGAL MEMATUHI / SECTION 5: CONSEQUENCES OF FAILURE TO COMPLY", 1)
    add_bilingual_block(doc, "This Performance Improvement Plan is a formal part of the disciplinary... process. Failure to meet the... objectives... may result in... termination of employment...", "Pelan Peningkatan Prestasi ini adalah sebahagian daripada proses tatatertib... Kegagalan untuk mencapai objektif... boleh mengakibatkan... penamatan perkhidmatan...")
    doc.add_heading("BAHAGIAN 6: AKUAN / SECTION 6: ACKNOWLEDGEMENT", 1)
    doc.add_paragraph("Akuan Pekerja / Employee Acknowledgement:")
    add_bilingual_block(doc, "I acknowledge that I have received a copy of this PIP... I understand the performance gaps, the objectives, and the consequences...", "Saya mengaku bahawa saya telah menerima salinan PIP ini... Saya faham akan jurang prestasi, objektif-objektif, dan akibat...")
    doc.add_paragraph("\n\n___________________\nTandatangan / Signature:\nNama / Name:\nTarikh / Date:\n")
    doc.add_paragraph("Akuan Pengurus / Manager Acknowledgement:")
    add_bilingual_block(doc, "I have discussed this plan with the employee and commit to providing the support...", "Saya telah membincangkan pelan ini dengan pekerja dan komited untuk menyediakan sokongan...")
    doc.add_paragraph("\n\n___________________\nTandatangan / Signature:\nNama / Name:\nTarikh / Date:")
    add_footer(doc, BRAND_TAGLINE); doc.save(os.path.join(save_folder, "05_Performance_Improvement_Plan_Template.docx"))
    print("...Performance Toolkit DONE.")
    return True # Indicate success

def make_documentation(COMPANY_DETAILS, BRAND_TAGLINE, logo_path, save_folder):
    """Generates the User Guide PDF and the README.txt file."""
    print("Generating Documentation...")
    
    # === User_Guide.pdf ===
    styles = getSampleStyleSheet()
    guide_pdf = os.path.join(save_folder, "User_Guide.pdf")
    story = []
    try:
        if logo_path and os.path.exists(logo_path):
            # Use the sanitized logo path
            story.append(Image(logo_path, width=Inches(1), height=Inches(1)))
    except Exception as e:
        print(f"⚠️ PDF Logo add failed: {e}")
        # This error is now critical, as a bad logo would have been caught by the sanitizer
        # But we can still proceed without it.
        
    story += [
        Spacer(1, 12),
        Paragraph(f"{COMPANY_DETAILS['name']} — HR & People Management Pack", styles['Title']),
        Paragraph(BRAND_TAGLINE, styles['Italic']),
        Spacer(1, 12),
        Paragraph("<b>Pack Contents:</b>", styles['h3']),
        Paragraph("<b>1. Hiring & Onboarding Kit:</b> Job Descriptions, Checklists, Interview Forms, Offer Letter, Welcome Letter, and Onboarding Plan.", styles['Normal']),
        Paragraph("<b>2. Employee Handbook:</b> A comprehensive, compliant handbook template covering all major policies from the Employment Act 1955.", styles['Normal']),
        Paragraph("<b>3. Performance Management Toolkit:</b> Performance Review, Self-Evaluation, SMART/OKR Goals, 1-on-1, and PIP templates.", styles['Normal']),
        Spacer(1, 12),
        Paragraph("<b>Instructions:</b>", styles['h3']),
        Paragraph(f"1. This pack has been pre-filled with your company details (e.g., <b>{COMPANY_DETAILS['name']}</b>) and branded with your logo.", styles['Normal']),
        Paragraph("2. Open the `.docx` files in Microsoft Word or Google Docs.", styles['Normal']),
        Paragraph("3. Find and replace all remaining employee-specific placeholders (e.g., <b>&lt;&lt;Nama Calon&gt;&gt;</b>, <b>&lt;&lt;Jawatan / Position&gt;&gt;</b>) with the new hire's details.", styles['Normal']),
        Paragraph("4. Review all clauses, especially in the Offer Letter and Handbook, to ensure they match your company's policies.", styles['Normal']),
        Spacer(1, 12),
        Paragraph("<i>Disclaimer: These templates are not legal advice. Always consult a qualified legal or HR professional before implementation.</i>", styles['Italic']),
    ]
    
    try:
        SimpleDocTemplate(guide_pdf).build(story)
    except Exception as e:
        print(f"CRITICAL ERROR: PDF generation failed: {e}")
        # This could happen if the logo is *still* bad, though unlikely
        raise Exception(f"PDF generation failed. Logo may be invalid. Error: {e}")

    # === README.txt ===
    readme_content = f"""{COMPANY_DETAILS['name']}
{BRAND_TAGLINE}

=======================================================
[EN] Welcome to the Malaysian SME HR & People Management Template Pack
=======================================================

Thank you for choosing this template pack. This resource is designed to provide Malaysian Small and Medium Enterprises (SMEs) with a practical, compliant, and professional set of HR documents.

[IMPORTANT: READ THIS FIRST]
This document is a template and not legal advice. The documents are based on the principles of the Malaysian Employment Act 1955 (including 2023 amendments), the Industrial Relations Act 1967, and common HR best practices.

The legal landscape is complex; these templates are a starting point, not a complete solution.

Your Responsibility:
1. Customise: You must find and replace all placeholders (e.g., <<Employee Name>>, <<Position>>). Company-level details (like '{COMPANY_DETAILS['name']}') have been pre-filled.
2. Review: You must review every clause to ensure it matches your company's actual policies (e.g., specific leave days above the statutory minimum, company-specific benefits, dress code).
3. Seek Legal Advice: Before implementing these documents, we strongly recommend you have them reviewed by a qualified Malaysian legal or HR compliance professional to ensure they are perfectly tailored to your specific business needs.

We hope this pack helps you build a great, productive, and compliant workplace.

=======================================================
[BM] Selamat Datang ke Pek Templat HR & Pengurusan Staf PKS Malaysia
=======================================================

Terima kasih kerana memilih pek templat ini. Sumber ini direka untuk menyediakan PKS (Perusahaan Kecil dan Sederhana) Malaysia dengan satu set dokumen HR yang praktikal, patuh undang-undang, dan profesional.

[PENTING: BACA DAHULU]
Dokumen ini hanyalah templat dan bukan nasihat undang-undang. Dokumen-dokumen ini disediakan berdasarkan prinsip Akta Kerja 1955 (termasuk pindaan 2023), Akta Perhubungan Perusahaan 1967, and amalan terbaik HR semasa.

Persekitaran undang-undang adalah rumit; templat ini adalah titik permulaan, bukan penyelesaian muktamad.

Tanggungjawab Anda:
1. Suaikan: Anda mesti mencari dan menggantikan semua pemegang tempat (cth., <<Nama Pekerja>>, <<Jawatan>>). Butiran peringkat syarikat (seperti '{COMPANY_DETAILS['name']}') telah diisi terlebih dahulu.
2. Semak: Anda mesti menyemak setiap fasal untuk memastikannya sepadan dengan polisi sebenar syarikat anda (cth., bilangan cuti tambahan melebihi had minimum statutori, faedah khusus syarikat, etika pakaian).
3. Dapatkan Nasihat Guaman: Sebelum melaksanakan dokumen-dokumen ini, kami amat mengesyorkan agar ia disemak oleh pakar undang-undang atau pakar pematuhan HR Malaysia yang bertauliah untuk memastikan ia disesuaikan dengan sempurna untuk keperluan khusus perniagaan anda.

Kami berharap pek ini membantu anda membina tempat kerja yang hebat, produktif, dan patuh undang-undang.
"""
    with open(os.path.join(save_folder, "README.txt"),"w",encoding="utf-8") as f:
        f.write(readme_content)
    print("...Documentation DONE.")
    return True # Indicate success


# ===============================================================
# 5️⃣ The Master Function (This is what the API calls)
#
# UPDATED: Now calls sanitize_logo()
# ===============================================================

def generate_hr_pack(COMPANY_DETAILS, BRAND_TAGLINE, logo_upload_file):
    
    print(f"Starting pack generation for: {COMPANY_DETAILS['name']}")
    
    # --- 1. Create a temporary, unique directory to build the pack
    base_dir = tempfile.mkdtemp()
    
    # Define all the folder paths
    pack_root = os.path.join(base_dir, "HR_People_Management_Pack")
    kit_folder = os.path.join(pack_root, "Hiring_Onboarding_Kit")
    handbook_folder = os.path.join(pack_root, "Employee_Handbook")
    performance_folder = os.path.join(pack_root, "Performance_Management_Toolkit")
    doc_folder = os.path.join(pack_root, "Documentation")
    
    # Create the directory structure
    for f in [kit_folder, handbook_folder, performance_folder, doc_folder]:
        os.makedirs(f, exist_ok=True)

    # --- 2. Process and Sanitize the uploaded logo
    sanitized_logo_path = None
    if logo_upload_file:
        print(f"Sanitizing logo: {logo_upload_file.filename}")
        sanitized_logo_path = sanitize_logo(logo_upload_file, base_dir)
        # Note: sanitize_logo will raise an exception if it fails

    # --- 3. Run all the generator functions
    # All functions now receive the path to the NEW sanitized logo
    try:
        make_hiring_kit(COMPANY_DETAILS, BRAND_TAGLINE, sanitized_logo_path, kit_folder)
        make_handbook(COMPANY_DETAILS, BRAND_TAGLINE, sanitized_logo_path, handbook_folder)
        make_performance(COMPANY_DETAILS, BRAND_TAGLINE, sanitized_logo_path, performance_folder)
        make_documentation(COMPANY_DETAILS, BRAND_TAGLINE, sanitized_logo_path, doc_folder)
    except Exception as e:
        # Clean up the temp folder on error
        shutil.rmtree(base_dir)
        # Re-raise the exception to be caught by the API
        raise e

    # --- 4. Zip the entire generated folder
    print("Zipping the final pack...")
    safe_company_name = "".join(c for c in COMPANY_DETAILS['name'] if c.isalnum() or c in (' ', '_')).rstrip().replace(" ", "_")
    zip_filename = f"{safe_company_name}_HR_People_Management_Pack.zip"
    
    # We create the zip file in a place that's safe to clean up
    zip_output_path = os.path.join(base_dir, zip_filename)

    # Create the zip
    shutil.make_archive(
        base_name=os.path.join(base_dir, "HR_Pack_Final"), # A temporary base name
        format='zip',
        root_dir=base_dir,
        base_dir="HR_People_Management_Pack"
    )
    
    # Define the *final* zip file path
    final_zip_file = os.path.join(base_dir, "HR_Pack_Final.zip")

    print(f"Pack generated. Returning zip file: {final_zip_file}")

    # --- 6. Return the path to the zip file and its name
    # We can't clean up base_dir yet, because the zip file is in it.
    # The 'finally' block in the API route will handle cleanup.
    return final_zip_file, zip_filename, base_dir


# ===============================================================
# 6️⃣ The API Endpoint (The "Door")
#
# UPDATED: Now gets the logo file from 'request.files'
# ===============================================================

@app.route('/generate-pack', methods=['POST'])
def handle_generation():
    
    zip_path_to_send = None
    base_dir_to_clean = None
    
    try:
        # --- 1. Get data from the form ---
        
        # Get text data
        data = request.form
        COMPANY_DETAILS = {
            "name": data.get('company_name'),
            "address": data.get('company_address'),
            "working_hours": data.get('working_hours'),
            "dress_code": data.get('dress_code'),
            "ceo_name": data.get('ceo_name'),
            "ceo_title": data.get('ceo_title'),
            "hr_name": data.get('hr_name'),
            "hr_title": data.get('hr_title')
        }
        BRAND_TAGLINE = data.get('brand_tagline')
        
        # Get file data
        logo_upload_file = request.files.get('logo_upload')
        
        # Basic validation
        if not COMPANY_DETAILS['name']:
            raise Exception("Company Name is a required field.")

        # --- 2. Run the master generator function ---
        zip_path_to_send, zip_filename, base_dir_to_clean = generate_hr_pack(
            COMPANY_DETAILS, 
            BRAND_TAGLINE, 
            logo_upload_file
        )

        # --- 3. Send the .zip file back to the user ---
        return send_file(
            zip_path_to_send,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )

    except Exception as e:
        print(f"--- ERROR: {e} ---")
        # Send a JSON error back to the frontend
        return jsonify(error=f"An error occurred. {e}"), 500
    
    finally:
        # --- 4. Clean up the temporary directory ---
        # This block runs *after* the file is sent (or if an error occurs)
        if base_dir_to_clean and os.path.exists(base_dir_to_clean):
            try:
                shutil.rmtree(base_dir_to_clean)
                print(f"Successfully cleaned up temp directory: {base_dir_to_clean}")
            except Exception as e:
                print(f"Warning: Failed to clean up temp directory {base_dir_to_clean}. Error: {e}")

# Homepage Route (for testing)
@app.route('/')
def index():
    return jsonify(message="HR Pack Generator API is live and running."), 200

# --- 7. Run the App ---
if __name__ == "__main__":
    app.run(debug=True, port=5000)
